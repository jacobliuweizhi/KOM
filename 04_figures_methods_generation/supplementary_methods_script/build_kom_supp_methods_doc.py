from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_DIR = Path(__file__).resolve().parent
MD_PATH = OUT_DIR / "KOM_补充方法学完整版_审稿版_v2深度版_20260615.md"
DOCX_PATH = OUT_DIR / "KOM_补充方法学完整版_审稿版_v2深度版_20260615.docx"
CHECK_PATH = OUT_DIR / "KOM_补充方法学_覆盖性自查表_v2深度版_20260615.md"


CONTENT = r'''# KOM 膝骨关节炎医生端决策支持系统：补充方法学完整版（审稿版）

版本：2026-06-15  
用途：用于主文 Supplementary Methods。本文按审稿可复核逻辑补全标准化病例、KOM-Profile、KOM-Rad/OAK-Net、QM-Risk/KOM-Risk、SHAP、KOM-KB、KOM-RAG、KOM-MDT、KOM-Safe、KOM-Score、KOM-Sim 医生交互和消融实验的完整方法。  
写作边界：本文将“已锁定结果”“本地实现证据”“当前缺口”分开陈述。没有恢复到原始文件或 locked package 的信息，不写成已完成结果。

## S1. 研究设计、版本锁定和总体工作流

本研究为一个医生端膝骨关节炎（knee osteoarthritis, KOA）临床决策支持系统的开发、标准化病例验证和医生交互模拟研究。KOM 系统被设计为面向医生的辅助工作流，而非自动处方系统。系统将标准化病例资料转化为结构化患者画像，结合影像结构评估、纵向风险提示、指南锚定证据检索、多专科智能体协商、结构化处方生成和独立安全审计，最终形成医生可阅读、可修改、可追踪的治疗建议。

本文方法学遵循三个原则。第一，所有模型、检索和医生交互评价均以预先定义的输入、输出、评价指标和错误分类为核心，避免事后选择性报告。第二，涉及临床 AI 的报告尽量对齐 TRIPOD+AI、DECIDE-AI、CONSORT-AI 和 CLAIM 2024 等透明报告原则；其中 KOM-Risk 属于预测模型报告范围，KOM-Rad/OAK-Net 属于医学影像 AI 报告范围，KOM-Sim 属于早期医生端临床决策支持评价。第三，所有可复核结果必须能追溯至 locked package、front locked sheets、原始医生任务记录、专家评分表、RAG query/top-k/gold-label 输出或本地验证报告。

[图 S1 占位：KOM 总体方法流程图。建议流程为：OAI/标准化病例输入 -> KOM-Profile -> KOM-Rad/OAK-Net -> QM-Risk/KOM-Risk -> KOM-KB/KOM-RAG -> KOM-MDT -> KOM-Rx -> KOM-Safe -> KOM-Score/KOM-Sim 评价。]

## S2. 120 例标准化病例构建

### S2.1 数据来源与病例单位

标准化病例来自 Osteoarthritis Initiative（OAI）候选病例池。OAI 是一个多中心、纵向、前瞻性 KOA 观察队列，提供膝关节影像、症状、功能、体重、合并症、用药、治疗和随访资料。标准化病例的目标不是估计真实世界患病率，而是构建一组能覆盖医生端 KOA 治疗决策场景的研究任务集。

每一例标准化病例对应一个去标识化患者任务单元。病例构建时需保留可追溯的 OAI 受试者 ID、膝侧、访视时间和源变量索引，但在医生实验和评分阶段隐藏可识别信息。若同一受试者存在双膝资料，病例构建和后续任务分配必须避免把同一受试者的两个膝关节作为相互独立的训练或验证证据来解释；在风险建模中，同一受试者双膝必须保持在同一数据 split 中。

### S2.2 候选病例初筛

候选病例首先按核心字段完整性筛选。进入候选池的病例需要具备以下信息中的主要部分：人口学信息、身高体重或 BMI、膝侧、基线影像或影像结构标签、疼痛或功能指标、至少一个治疗决策相关变量、用药风险或合并症信息、治疗目标或活动需求信息。明显缺少病例构建核心信息者不进入标准化病例候选池。

初筛不是随机抽样，而是面向医生决策任务的病例工程。其目的是让后续 KOM-Profile、KOM-RAG、KOM-MDT 和医生处方任务能够面对不同临床形态，而不是重复同一类低信息量患者。

### S2.3 多维覆盖矩阵，而非四象限

本版方法不再使用“四象限”作为论文中的标准化病例构建框架。标准化病例改用“多维覆盖矩阵 + 分层目的性抽样 + 最大变异抽样”表述。早期内部标签若仍存在，仅作为历史审计标签，不作为主文病例选择结构、模型输入或 subgroup claim。

多维覆盖矩阵包含六类轴：

- 疾病结构轴：KL 分级、关节间隙狭窄、骨赘、硬化、单侧或双侧受累。
- 症状功能轴：疼痛强度、静息痛或活动痛、功能限制、KOOS/WOMAC 或等价功能字段。
- 代谢负荷轴：BMI、体重管理需求、肥胖相关机械负荷。
- 安全风险轴：胃肠、肾脏、心血管、跌倒、用药禁忌和注射/手术边界。
- 管理需求轴：职业或运动负荷、治疗期望、依从性风险、共同决策复杂度、随访需求。
- 治疗场景轴：运动康复、体重营养、镇痛/NSAIDs、注射、支具/物理治疗、手术转诊和长期随访。

### S2.4 抽样策略

病例构建采用分层目的性抽样和最大变异抽样。第一步从 OAI 候选病例池建立可构建病例清单；第二步计算病例在六类覆盖轴上的标签；第三步优先选择在结构严重度、症状负担、代谢负荷、安全禁忌和管理需求上互补的病例；第四步检查是否存在重复患者、重复场景或过度相似病例；第五步将病例整理为统一输入模板。最终锁定 120 例标准化病例，用于系统开发、KOM-Profile 抽取验证、RAG query 构建、KOM-MDT 处方生成、系统消融和医生交互任务。

### S2.5 证明病例选择符合原设计的可审计证据

病例构建完成后应输出以下证据，用于证明 120 例病例不是同一患者或同质患者的重复，而是覆盖多种 KOA 管理场景：

| 证据文件或检查 | 审计目的 | 通过标准 |
|---|---|---|
| 候选病例 manifest | 证明病例来自预定义候选池 | 每例有源 ID、膝侧、访视和源变量索引 |
| 去重检查 | 证明不是同一患者重复 | 去标识化患者 ID 无不当重复；双膝规则有说明 |
| 字段完整性表 | 证明病例可用于系统输入 | 核心字段存在率和缺失字段记录可复核 |
| 多维覆盖矩阵 | 证明病例具有代表性 | 结构、症状、代谢、安全、管理需求、治疗场景均有覆盖 |
| 病例模板锁定表 | 证明医生和系统看到同一病例输入 | 系统臂和医生臂使用同一标准化病例文本 |
| 版本冻结记录 | 证明后续实验未修改病例 | 病例文本、任务 ID、版本号和修改日志锁定 |

[图 S2 占位：120 例标准化病例构建流程图。建议画成“候选池 -> 完整性初筛 -> 多维覆盖矩阵 -> 最大变异抽样 -> 去重/审计 -> 120例病例锁定”。]

## S3. KOM-Profile 患者画像自动抽取

### S3.1 模块定位

KOM-Profile 是 KOM-Assess 的入口模块。其任务是将标准化病例文本、OAI 源变量和规则派生变量转化为结构化患者画像。患者画像供后续 KOM-RAG query 构建、KOM-MDT 协商、KOM-Rx 处方生成和 KOM-Safe 安全审计使用。KOM-Profile 不直接生成治疗建议，其作用是把病例资料变成可被系统和医生共同解释的结构化状态。

### S3.2 输入、字段体系和输出

KOM-Profile 输入包括三类信息：一是标准化病例文本；二是 OAI 源变量或源变量索引；三是规则派生变量。输出为一个结构化 profile，包括身份与基本信息、症状负担、功能限制、影像结构信息、合并症与用药安全风险、治疗目标、活动需求、关键风险和缺失字段提示。

字段分为三类。第一类为源真值可核对字段，例如年龄、性别、BMI、KL 分级、疼痛评分和 KOOS/WOMAC 等。第二类为规则派生字段，例如体重管理需求、药物安全风险、结构负担标签和管理需求标签。第三类为自由文本或不稳定核对字段，例如运动习惯、依从性预期、平衡功能和治疗偏好。

### S3.3 构建与“训练”方案

KOM-Profile 采用 schema-first 的规则化抽取流程，而不是以黑箱监督学习模型作为主体。开发过程包括：定义字段字典；为每个字段指定源变量、允许取值、单位、缺失值编码和临床解释；将源变量映射为标准字段；对规则派生字段编写确定性规则；将抽取结果与病例文本进行一致性检查；最后输出患者画像和字段来源表。

若系统在文本组织或字段补全中使用语言模型，其作用限定为格式化、摘要和字段候选生成；可核对字段必须回到源变量或病例文本核验。模型不得凭空补充病例中不存在的信息。每个 profile 输出均应保留字段来源、抽取规则、是否可核对和是否用于下游处方。

### S3.4 评价指标

字段级评价采用 exact match、partial match、wrong/missing 和 not_available 四类。Exact match 表示字段值与源真值一致；partial match 表示方向或主要临床含义正确但粒度不完全一致；wrong/missing 表示错误、遗漏或不可解释；not_available 表示源资料本身不支持该字段。对可核对字段计算字段级准确率、F1、完全匹配数、部分正确数和失败模式分布。

当前草案锁定结果显示，KOM-Profile 总字段框架包含 124 个字段，其中 56 个字段为源真值可核对字段，68 个字段为规则派生、管理性或不可稳定核对字段。在 56 个可核对字段中，字段级准确率/F1 为 0.846，31 个完全匹配，25 个部分正确。失败主要集中在运动习惯、体重管理需求、静息痛、下肢肌力、治疗依从性预期和平衡功能。论文中建议同时报告“全字段框架”和“可核对字段分母”，避免把规则派生字段错误纳入 source-truth 分母。

[图 S3 占位：KOM-Profile 字段抽取与核验流程图。建议包含：病例文本/OAI变量 -> 字段字典 -> 源字段抽取 -> 规则派生 -> 缺失标记 -> 字段级审计 -> 患者画像输出。]

## S4. KOM-Rad / OAK-Net 影像分析模块

### S4.1 模块定位

KOM-Rad 是 KOM-Assess 中的影像结构支持模块，内部算法为 OAK-Net。该模块用于从膝关节 X 线影像中提供结构严重度和不确定性信息，增强患者画像和风险解释。KOM-Rad 的输出不直接决定治疗处方，也不替代影像科或骨科医生判断。

### S4.2 输入与参考标准

输入为 OAI 膝关节 X 线影像及其对应的中心化半定量读片标签。参考标准包括 KL 分级和结构征象，例如关节间隙狭窄、骨赘和硬化等。所有影像样本在进入训练前应完成膝侧匹配、图像质量检查、标签完整性检查和患者级数据隔离。报告影像 AI 方法时应对齐 CLAIM 2024，说明数据来源、纳入排除标准、参考标准、预处理、模型架构、训练方案、验证策略和失败模式。

### S4.3 架构设计

主模型为 ConvNeXt-B backbone 加 evidential uncertainty head。ConvNeXt-B 用于提取 X 线影像的局部和层级结构特征；不确定性 head 用于输出预测可信度或证据强度，辅助识别低置信度样本。DenseNet-121 作为轻量级对照模型，用于评估主模型相对传统卷积架构的增益和稳定性。

模型输出包括 KL 各等级概率、预测 KL 等级、结构征象提示和不确定性分数。若不确定性超过预设阈值，KOM-Rad 输出应进入“需医生复核”状态，而不是作为高置信度自动结论进入处方。

### S4.4 训练和实验设计

影像模块训练应采用患者级或受试者级 split，确保同一患者影像不跨训练、验证和测试集。预处理包括图像方向统一、膝关节区域裁剪或定位、灰度归一化、分辨率标准化和质量控制。训练阶段可使用随机裁剪、亮度/对比度扰动、小幅旋转等不会改变结构标签的增强策略。模型选择以验证集表现为准，测试集只用于最终报告。

评价指标包括 quadratic weighted kappa（QWK）、balanced accuracy（BACC）、macro-F1、mean absolute error（MAE）、expected calibration error（ECE）和 selective accuracy at 80% coverage（sel_acc@80）。QWK 适合 KL 分级这种序数标签；BACC 和 macro-F1 处理类别不均衡；MAE 衡量序数偏差；ECE 衡量概率校准；sel_acc@80 衡量高置信度样本中的可靠性。

### S4.5 已锁定结果与当前边界

当前草案锁定结果为：ConvNeXt-B 达到 QWK 0.806 +/- 0.008、BACC 0.659、macro-F1 0.664、MAE 0.417、ECE 0.119、sel_acc@80 0.725。DenseNet-121 达到 QWK 0.805 +/- 0.007、BACC 0.669、macro-F1 0.664、MAE 0.420。恢复增强表未包含 image-level prediction，因此 image-level confusion matrix、calibration bins 和不确定性分布暂不能作为 locked 图直接绘制。

[图 S4 占位：影像分析模块流程图。建议包含：X线输入 -> 图像质控/裁剪 -> ConvNeXt-B/DenseNet -> KL概率/不确定性 -> 医生复核门控 -> KOM-Profile/KOM-Risk。]

## S5. QM-Risk / KOM-Risk PostDedup 纵向风险模块

### S5.1 模块定位

QM-Risk / KOM-Risk 是 KOM-Assess 中的纵向风险提示模块，用于估计 KOA 结构进展、固定时间窗内膝关节手术/TKR 事件和症状/功能恶化风险。该模块输出风险分层和解释，不生成治疗处方，不替代医生决策。预测模型报告应对齐 TRIPOD+AI。

### S5.2 输入特征和 endpoint

风险模块以 OAI 的膝关节水平数据为主要建模单位。输入特征包括人口学信息、BMI、基线症状和功能、影像结构严重度、合并症和安全风险、用药或治疗史、体力活动、心理行为相关变量和缺失指示变量。Endpoint A 为 KL structural progression；Endpoint B 为 TKR/knee surgery event，在非事件完整 censoring time 无法可靠重建时采用 fixed-horizon binary fallback；Endpoint C 为 symptom/function worsening，作为补充 endpoint。

Endpoint A 和 B 使用 56 个 raw features / 60 个 encoded features；Endpoint C 使用 60 个 raw features / 64 个 encoded features。所有特征字典、encoded feature names 和预处理追踪应保存到 locked package 中。

### S5.3 去重、侧别和泄漏控制

KXR SQ BU 文件族中 SIDE 映射经 wide right/left 字段交叉验证。草案锁定的映射检查显示 SIDE=1 到 right 的匹配率为 0.995270，SIDE=2 到 left 的匹配率为 0.996254。同一 ID-SIDE-visit 的多 READPRJ 记录按 15 > 37 > 42 > other 规则去重，并优先保留 KL 非缺失记录。

所有数据 split 必须在 person level 完成，同一受试者双膝不得跨 split。所有 preprocessing 只在训练集 fit；验证集用于模型和阈值选择；测试集只用于最终评价。基线模型不得纳入随访结局、未来影像、未来手术、未来用药升级或任何 endpoint proxy。

### S5.4 训练方案和候选模型

候选模型包括 elastic-net logistic regression、random forest、XGBoost、LightGBM 和 CatBoost。若 endpoint 以 survival 形式可靠构建，可另设 CoxPH、random survival forest 或 survival gradient boosting；若事件时间或 censoring 不能可靠重建，则使用 fixed-horizon binary endpoint。

最终模型选择不只依据 AUROC。模型必须同时满足判别能力、校准、临床净获益、风险分层可解释性、特征合理性和泄漏审计。二分类 endpoint 报告 AUROC、AUPRC、Brier score、校准曲线、ECE、DCA、敏感性、特异性、PPV、NPV 和 bootstrap CI。阈值应在验证集锁定，不能在测试集上重新优化。

### S5.5 输出和判定

每个 endpoint 最终保留一个 best_model.joblib，并保存 sample-level predictions、metrics、bootstrap CI、calibration、DCA、risk stratification、feature importance、SHAP 输出和 QC 文件。风险输出包括连续风险概率、预设阈值下的低/中/高风险分层、主要驱动特征和缺失/不确定性提示。具体阈值应从 locked package 导出；若阈值文件未恢复，论文中只报告连续风险和模型性能，不写具体分层阈值。

### S5.6 最新 locked 结果

| Endpoint | N rows | Persons | Events | Event rate | Best model | AUROC | AUPRC | Brier | Decision |
|---|---:|---:|---:|---:|---|---:|---:|---:|---|
| A: KL structural progression | 7,294 | 3,656 | 970 | 13.3% | CatBoost | 0.781 | 0.349 | 0.191 | ACCEPT_MAIN |
| B: TKR/knee surgery event | 9,592 | 4,796 | 548 | 5.7% | CatBoost | 0.868 | 0.362 | 0.128 | ACCEPT_MAIN; fixed-horizon binary |
| C: symptom/function worsening | 8,383 | 4,202 | 1,871 | 22.3% | CatBoost | 0.685 | 0.348 | 0.222 | ACCEPT_SUPPLEMENT |

[图 S5 占位：QM-Risk/KOM-Risk 建模流程图。建议包含：OAI膝关节数据 -> 侧别/去重 -> endpoint构建 -> person-level split -> 特征编码 -> 候选模型 -> CatBoost locked model -> 校准/DCA/SHAP -> 风险输出。]

## S6. SHAP 和模型解释

SHAP 用于解释 KOM-Risk 的预测驱动因素，而不是证明因果机制。解释输出分为全局解释和个体解释。全局解释包括 SHAP summary、平均绝对 SHAP 排名、top-feature plausibility audit 和特征方向检查；个体解释包括单例 force/waterfall 样式输出，用于显示某个患者风险升高或降低的主要变量。

SHAP 解释必须满足三条边界。第一，解释对象是 locked model 的 test 或 holdout predictions，不对训练集进行选择性展示。第二，解释只说明模型如何使用特征，不声称某个特征导致结局。第三，解释要结合临床合理性审计；若高贡献特征属于可能泄漏、未来信息或 endpoint proxy，该模型不得作为主模型接受。

## S7. KOM-KB 证据库构建

### S7.1 PubMed 检索策略

KOM-KB 的候选文献主要来自 PubMed 检索和指南/机构来源补充。2026-06-15 使用 NCBI E-utilities 复核的 PubMed 候选检索式如下：

```text
("knee osteoarthritis"[tiab] OR "osteoarthritis of the knee"[tiab] OR gonarthrosis[tiab])
AND
(management[tiab] OR treatment[tiab] OR intervention[tiab] OR guideline[pt]
OR practice guideline[pt] OR systematic review[pt] OR meta-analysis[pt]
OR randomized controlled trial[pt] OR exercise[tiab] OR rehabilitation[tiab]
OR "weight loss"[tiab] OR NSAID[tiab] OR injection[tiab]
OR "physical therapy"[tiab] OR pain[tiab])
AND humans[MeSH Terms] AND english[lang]
AND ("2018/01/01"[Date - Publication] : "2026/06/15"[Date - Publication])
```

该检索式在 2026-06-15 返回 7,282 条 PubMed 候选记录。该数量代表候选检索集，不等于最终 evidence unit 数。Evidence unit 是文献经过去重、筛选、分层、拆分和结构化后的最小证据单位。

### S7.2 Evidence unit 定义

KOM-KB 的最小单位是 evidence unit，即一个可被检索、路由、引用和审计的证据记录。每个 evidence unit 至少包含：唯一 ID、标题、来源、年份、文献类型、证据等级、治疗域、适用人群、干预或暴露、对照、结局、安全标签、专科归属、证据摘要、适用/不适用边界、原始链接和审计状态。

### S7.3 L1-L7 证据层级

为满足审稿可解释性，建议使用七级证据层级。当前草案 locked count 报告为五个压缩层级，因此七级未压缩 count 需要从 evidence database 重新导出，不能凭空拆分。

| 层级 | 定义 | 用途 |
|---|---|---|
| L1 | 指南、临床标准、共识推荐、监管或机构级安全边界 | 作为 KOM-RAG 和 KOM-Safe 的最高锚点 |
| L2 | 系统综述、meta-analysis、umbrella review、证据综合 | 支持治疗域总体有效性和安全性 |
| L3 | 随机对照试验、前瞻性干预研究 | 支持具体干预建议 |
| L4 | 观察性研究、预后/诊断/真实世界证据 | 支持风险、适用性和背景判断 |
| L5 | 药物安全、禁忌证、剂量/用药边界和不良反应证据 | 支持 KOM-Safe 门控 |
| L6 | 康复方案、实施路径、患者教育、依从性和随访资料 | 支持可执行处方细节 |
| L7 | 背景、机制、专家评论或低优先级资料 | 仅作背景，不单独驱动强建议 |

草案锁定的压缩证据库包含 3,266 个 evidence units 和 2,174 个唯一来源：L1 指南/临床标准 99 条、L2 证据综合 648 条、L3 干预性临床证据 1,124 条、L4 观察/预后证据 488 条、L5 背景/方案/实施证据 907 条。若最终主文坚持 L1-L7，应重新导出七级 counts。

[图 S6 占位：KOM-KB evidence unit 构建图。建议包含：PubMed候选文献/指南来源 -> 去重 -> 题录筛选 -> 证据拆分 -> L1-L7分级 -> 标签体系 -> evidence unit入库。]

## S8. KOM-RAG / GraphRAG 检索和评价

### S8.1 检索输入和返回流程

KOM-RAG 接收来自 KOM-Profile 的患者画像、治疗域需求和安全风险标签。Query 包含结构严重度、疼痛/功能负担、BMI 或体重管理需求、用药安全风险、活动目标、注射/手术边界和随访需求。系统先进行治疗域路由，再在 KOM-KB 中检索候选 evidence units，并根据指南锚点、证据层级、适用性、安全标签、专科标签和排序分数返回 top-k 证据。

返回结果分为 direct evidence、context evidence 和 safety evidence。Direct evidence 直接支持某一治疗建议；context evidence 支持患者背景、风险或适用性判断；safety evidence 用于限制、降级或否决某个建议。KOM-MDT 和 KOM-Safe 只能使用可追溯 evidence unit，不得引用无法定位的泛化知识。

### S8.2 Naive RAG baseline

Naive RAG baseline 使用同一批 query 和同一证据库，但仅执行单阶段向量相似度 top-k 检索。它不使用指南锚点、证据层级、专科路由、安全标签、图结构关系或证据仲裁。该 baseline 用于回答：KOM-RAG 的结构化证据路由是否优于普通检索增强。

### S8.3 金标准构建

RAG 评价的 gold standard 以 query-level manual review pack 为基础。每个 query 的 gold evidence 由人工审阅候选 evidence units，标注 relevant、partially relevant、not relevant 和 unsafe/misapplied。推荐流程为两名独立标注者初评，分歧由高年资临床专家或方法学负责人仲裁；如果当前文件只恢复到 gold-detail/manual review pack 而未恢复标注者数量，则正文只能写“人工审阅和仲裁形成 gold labels”，不能写具体标注者数量。

### S8.4 指标计算

Precision@10 衡量前 10 条返回中 relevant evidence 的比例。Recall@K 衡量 gold evidence 在前 K 条中被召回的比例。Hit@10 衡量前 10 条是否至少命中 1 条 relevant evidence。MRR 衡量第一个 relevant evidence 的倒数排名。nDCG@10 衡量带等级相关性的排序质量。所有指标均在同一 query benchmark 上分别计算 KOM-RAG 和 naive RAG baseline。

### S8.5 锁定结果和缺口

锁定 benchmark 包含 160 个 query。KOM-RAG Precision@10 = 0.676，naive RAG = 0.303；MRR = 0.748 vs 0.159；nDCG@10 = 0.690 vs 0.237；Hit@10 = 1.000 vs 0.688。KOM-RAG Recall@10/20/27/30 = 0.412/0.695/0.824/0.855。

当前恢复数据支持检索级评价，但 graph node/edge manifest、embedding model、reranker config、faithfulness、citation support、unsupported claim rate 和 safety evidence hit rate 尚未恢复。因此主文不能声称 generation grounding 或 citation faithfulness 已完成；可作为待补跑项列入补充材料。

[图 S7 占位：KOM-RAG 与 naive RAG 对照流程图。建议显示同一query和同一证据库下，baseline只做向量top-k，KOM-RAG加入指南锚点/证据层级/安全标签/专科路由/仲裁。]

## S9. KOM-MDT 多专科智能体协商

### S9.1 总体流程

KOM-MDT 将患者画像和 KOM-RAG 证据输入多个专科视角，按“初稿-交叉质疑-证据仲裁-骨科仲裁-最终处方”的顺序运行。其设计目标不是模拟真实医院 MDT 的人员构成，而是把 KOA 处方中常见的运动、体重营养、心理行为、骨科边界和证据安全问题显式拆开，减少单一生成模型遗漏安全门控或治疗组件的风险。

### S9.2 R0-R8 智能体角色

| 阶段 | 智能体 | 输入 | 输出 |
|---|---|---|---|
| R0 | 患者画像共识 | KOM-Profile、KOM-Rad、KOM-Risk、缺失字段 | 症状、结构、安全风险、治疗目标和禁忌的共识摘要 |
| R1 | 运动医学智能体 | 画像、运动目标、疼痛/结构负担、exercise evidence | FITT 运动处方、强度进阶、停止规则 |
| R2 | 体重营养代谢智能体 | BMI、体重目标、代谢风险、nutrition evidence | 体重目标、营养建议、肌肉保护策略 |
| R3 | 心理行为智能体 | 依从性、疼痛认知、生活方式和随访需求 | 行为管理、教育、依从性和随访建议 |
| R4 | 骨科综合智能体 | 影像结构、症状、注射/手术边界 | 药物、注射、转诊、保守治疗边界 |
| R5 | 跨专科质疑智能体 | R1-R4 草案 | 冲突、遗漏、禁忌和证据不足清单 |
| R6 | 证据仲裁智能体 | RAG证据、指南锚点、R5质疑 | 保留、降级、删除或要求补证的建议 |
| R7 | 骨科仲裁智能体 | 结构严重度、注射/手术边界、R6结果 | 最终骨科边界和升级条件 |
| R8 | 最终处方智能体 | 仲裁后建议、KOM-Safe结果 | 医生可读 KOM-Rx 和安全审计摘要 |

### S9.3 Prompt 和知识源构建

每个智能体 prompt 应由四部分组成：角色边界、输入字段、证据使用规则和输出模板。角色边界规定智能体只能处理本专科问题；输入字段来自 KOM-Profile 和 RAG 证据；证据使用规则要求引用 evidence unit ID 和证据等级；输出模板要求列出建议、适用条件、禁忌、证据 ID 和需医生复核项。Prompt 不应允许智能体越权诊断、编造证据或绕过安全审计。

知识源由 KOM-KB evidence units、指南锚点、安全标签和专科标签组成。各智能体不是访问不同的“私有知识库”，而是在同一证据库上通过专科标签和治疗域路由获得不同证据视角。这样可保证证据来源一致，同时允许不同专科对同一病例提出不同关注点。

[图 S8 占位：KOM-MDT R0-R8 协商流程图。建议画为纵向审计链，突出 R5 质疑、R6 证据仲裁、R7 骨科仲裁和 R8 安全释放。]

## S10. KOM-Safe / KVMSafe 安全审计

KOM-Safe（用户提到的 KVMSafe 可在文中统一命名为 KOM-Safe 或在首次出现时写作 KOM-Safe/KVMSafe）是独立安全审计模块。它不生成新的治疗建议，而是检查 KOM-MDT 和 KOM-Rx 输出是否违反安全规则、指南边界或病例禁忌。

安全审计覆盖七类规则域：NSAIDs 和镇痛药禁忌；胃肠、肾脏、心血管和抗凝风险；注射适应证和禁忌；运动处方风险和停止规则；手术/转诊边界；红旗症状和升级条件；随访监测和风险复查。审计输出分为 pass、warning 和 critical。Warning 要求处方增加限制条件或医生复核；critical 要求阻断或重写相关建议。

KOM-Safe 与专科智能体的交互发生在两个时间点。第一，在 R5/R6 阶段，安全标签参与证据仲裁，用于降级或删除高风险建议。第二，在 R8 最终释放前，KOM-Safe 对完整 KOM-Rx 进行最终门控，生成安全问题、修改建议、错误分类和审计日志。任何被判定为 safety-critical 的错误均应进入 KOM-Score-Error。

## S11. KOM-Score 多源评价框架

### S11.1 组成

KOM-Score 不是单一评分，而是由 KOM-Score-Expert、KOM-Score-Rule 和 KOM-Score-Error 组成的多源评价框架。Expert 评分用于评价临床质量；Rule 评分用于复核结构化完整性和安全门控；Error 分类用于识别安全关键错误、临床决策错误和轻微完整性/表达错误。

### S11.2 Expert 评分

六名专家盲法评分由骨科、运动医学和康复医学各 2 名专家完成。专家隐藏处方来源和实验条件。评分维度包括总体质量、安全性、指南一致性、患者个体化、可执行性、证据可追溯性、专科完整性和临床一致性。各维度可使用 1-10 分或 0-100 分，最终统一换算到 0-100 分。专家评分需报告 ICC(2,1) 或相应一致性指标。

### S11.3 Rule 评分和 Error 分类

Rule 评分是确定性评分，不替代专家评价。规则域包括治疗维度完整性、运动处方参数、体重/营养、药物禁忌、注射边界、手术转诊、随访升级和 safety gate。Error 分类分为三层：safety-critical error、clinically relevant error 和 minor error。Safety-critical error 指可能导致患者实质性伤害、违反明确禁忌、错误建议或延误必要转诊、错误使用高风险药物，或遗漏关键安全门控的问题。

安全关键错误率计算为：

```text
Safety-critical error rate = number of safety-critical errors / number of prescription records x 100
```

草案锁定一致性结果显示，系统消融总体质量 ICC(2,1)=0.796，安全性 ICC=0.574；医生处方总体质量 ICC(2,1)=0.946，各维度 ICC=0.902-0.946。安全性 ICC 较低时，应结合规则评分和错误日志解释，而不单独依赖主观安全评分。

## S12. KOM-Sim 医生交互实验

### S12.1 实验对象和任务

KOM-Sim 是医生端标准化处方交互实验。最终纳入 26 名医生，每名医生完成 30 个标准化任务，形成 780 条医生-任务级处方记录。这里的 780 表示一名医生在一个任务中生成的一条处方记录，不代表 780 名医生或 780 个独立病例。

### S12.2 实验条件

实验包含四类条件。Clinician alone 条件下医生仅查看标准化病例资料；Clinician + KOM 条件下医生查看病例资料和 KOM 建议；Clinician + KOM-R 条件下医生查看病例资料、KOM 建议及 rationale/evidence；KOM standalone 是系统独立输出，用作系统基准，不作为医生操作臂混入。

### S12.3 记录变量

每条医生-任务记录应保存处方文本、处方来源条件、病例 ID、医生 ID、医生专业/年资/KOA 经验/AI 经验、编辑时间、治疗组件数、文本长度、KOM 建议查看次数、复制/采纳行为、工作负荷、信心、信息充分性、决策把握度、专家评分、规则评分、安全错误和高质量处方标记。历史旧表中存在 783/784 行 task-level sheet，最终主文只使用 780 条有效医生-任务记录，作图前必须使用 final_record_filter_log 过滤。

### S12.4 评价设计

主要比较为 Clinician + KOM vs Clinician alone，用于估计 KOM 建议对医生处方质量、规则完整性和安全错误的影响。次要比较为 Clinician + KOM-R vs Clinician + KOM，用于估计解释和证据展示的增量效应。KOM standalone 与医生条件分开解释，用于展示完整系统独立输出水平，而不声称其替代医生。

医生交互结果不得由系统自评作为唯一结论。处方质量应由盲法专家评分和规则评分共同评价；安全关键错误应由规则审计和人工复核共同确认。统计分析建议使用配对比较或混合效应模型，控制医生和病例的重复测量。

[图 S9 占位：KOM-Sim 医生交互实验流程图。建议包含：26名医生 -> 30个任务 -> 三个医生条件 + 系统独立基准 -> 处方记录 -> 盲法专家评分/规则评分/错误审计。]

## S13. KOM-MDT 和系统消融实验

系统消融用于评价完整 KOM 工作流中证据检索、MDT 协商和结构化工作流的贡献。当前 locked interface/report 展示四组：Full KOM、KOM without RAG、KOM without MDT 和 Direct LLM baseline。

| 组别 | 关闭或保留内容 | 评价目的 |
|---|---|---|
| Full KOM | 保留 KOM-Assess、KOM-RAG、KOM-MDT、KOM-Rx、KOM-Safe | 完整系统工作流 |
| KOM w/o RAG | 关闭外部证据检索，保留其他流程 | 评价证据路由贡献 |
| KOM w/o MDT | 关闭多专科协商和交叉质疑 | 评价 MDT 结构贡献 |
| Direct LLM | 不启用 KOM 结构，直接由通用生成流程输出 | 评价通用生成基线 |

消融输出使用同一组标准化病例、同一处方模板和同一 KOM-Score 评价。草案/本地 demo 显示消融结果为：Full KOM overall 8.46、安全性 9.11；KOM w/o RAG overall 6.56、安全性 7.99；KOM w/o MDT overall 6.44、安全性 7.91；Direct LLM overall 5.47、安全性 7.07。统计报告应补充配对检验、效应量和多重比较校正。

## S14. 统计分析

描述性统计根据变量类型报告均值/标准差、中位数/IQR、频数和比例。模型性能报告 AUROC、AUPRC、Brier、校准、DCA 和 bootstrap CI。RAG 检索报告 Precision@10、Recall@K、Hit@10、MRR 和 nDCG@10。医生交互实验优先使用配对设计或混合效应模型，医生和病例作为随机效应，实验条件作为固定效应。若数据分布不满足正态假设，可使用 Wilcoxon signed-rank test 或 permutation/sign-flip test。多重比较使用 Benjamini-Hochberg FDR 校正。

医生年资与 KOM 获益关系可用 LOWESS 可视化，并报告 Pearson 和 Spearman 相关。所有统计图需明确分母：医生数、病例数、医生-任务记录数、query 数或 evidence unit 数，避免把不同层级样本混用。

## S15. 缺失项和写作限制

| 缺失项 | 影响 | 写作处理 | 后续动作 |
|---|---|---|---|
| OAI 原始数据库本体未随文档上传 | 不能从 Excel 恢复原始数据 | 写明本地归档和 manifest，不重复分发原始 OAI | 保留 OAICompleteData_CSV/SAS 只读归档 |
| KOM-RAG graph node/edge manifest | 不能画真实 GraphRAG 网络结构 | 不写 node/edge 数 | 从 RAG 构建目录恢复 |
| RAG generation grounding | 不能写 faithfulness/citation support 已完成 | 仅写 retrieval-level evaluation | 后续补跑生成级评价 |
| OAK-Net image-level prediction | 不能画 image-level confusion/calibration bins | 只写 summary metrics | 从模型输出目录恢复 |
| KOM-Risk 阈值文件若未恢复 | 不能写具体低/中/高风险阈值 | 报告连续风险和性能 | 从 locked package 导出阈值表 |
| 七级 evidence counts 未导出 | 不能把 L1-L5 压缩 count 硬拆为 L1-L7 | 写七级定义，count 写压缩层级 | 重新导出 L1-L7 count |

## S16. 建议补充图位置

1. 图 S1：KOM 总体方法流程图。
2. 图 S2：120 例标准化病例构建流程图。
3. 图 S3：KOM-Profile 字段抽取与核验流程图。
4. 图 S4：KOM-Rad/OAK-Net 影像分析模块流程图。
5. 图 S5：QM-Risk/KOM-Risk 训练、验证、校准、DCA 和 SHAP 流程图。
6. 图 S6：KOM-KB evidence unit 构建和 L1-L7 分级流程图。
7. 图 S7：KOM-RAG vs naive RAG baseline 检索对照图。
8. 图 S8：KOM-MDT R0-R8 多专科协商流程图。
9. 图 S9：KOM-Sim 医生交互实验和 KOM-Score 评价流程图。

## S17. 参考文献和方法学依据

1. Osteoarthritis Initiative. Osteoarthritis Initiative: a knee health study. NIAMS/NDA. https://nda.nih.gov/oai and https://www.niams.nih.gov/grants-funding/funded-research/osteoarthritis-initiative.
2. Vasey, B. et al. Reporting guideline for the early-stage clinical evaluation of decision support systems driven by artificial intelligence: DECIDE-AI. Nat. Med. 28, 924-933 (2022).
3. Liu, X. et al. Reporting guidelines for clinical trial reports for interventions involving artificial intelligence: the CONSORT-AI extension. Nat. Med. 26, 1364-1374 (2020).
4. Collins, G. S. et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ 385, e078378 (2024).
5. Tejani, A. S. et al. Checklist for Artificial Intelligence in Medical Imaging (CLAIM): 2024 update. Radiol. Artif. Intell. 6, e240300 (2024).
6. Kolasinski, S. L. et al. 2019 American College of Rheumatology/Arthritis Foundation guideline for the management of osteoarthritis of the hand, hip, and knee. Arthritis Rheumatol. 72, 220-233 (2020).
7. Bannuru, R. R. et al. OARSI guidelines for the non-surgical management of knee, hip, and polyarticular osteoarthritis. Osteoarthritis Cartilage 27, 1578-1589 (2019).
8. National Institute for Health and Care Excellence. Osteoarthritis in over 16s: diagnosis and management. NICE guideline NG226 (2022).
9. Brophy, R. H. & Fillingham, Y. A. AAOS clinical practice guideline summary: management of osteoarthritis of the knee (nonarthroplasty), third edition. J. Am. Acad. Orthop. Surg. 30, e721-e729 (2022).
10. Bruyere, O. et al. An updated algorithm recommendation for the management of knee osteoarthritis from ESCEO. Semin. Arthritis Rheum. 49, 337-350 (2019).
11. Moseng, T. et al. EULAR recommendations for the non-pharmacological core management of hip and knee osteoarthritis: 2023 update. Ann. Rheum. Dis. 83, 730-740 (2024).
12. Lewis, P. et al. Retrieval-augmented generation for knowledge-intensive NLP tasks. Adv. Neural Inf. Process. Syst. 33, 9459-9474 (2020).
13. Edge, D. et al. From local to global: a graph RAG approach to query-focused summarization. Preprint at https://arxiv.org/abs/2404.16130 (2024).
14. Lundberg, S. M. & Lee, S.-I. A unified approach to interpreting model predictions. Adv. Neural Inf. Process. Syst. 30, 4765-4774 (2017).
15. Guo, C. et al. On calibration of modern neural networks. Proc. 34th Int. Conf. Mach. Learn. 70, 1321-1330 (2017).
16. Brier, G. W. Verification of forecasts expressed in terms of probability. Mon. Weather Rev. 78, 1-3 (1950).
17. Vickers, A. J. & Elkin, E. B. Decision curve analysis: a novel method for evaluating prediction models. Med. Decis. Making 26, 565-574 (2006).
18. Cohen, J. Weighted kappa: nominal scale agreement with provision for scaled disagreement or partial credit. Psychol. Bull. 70, 213-220 (1968).
19. Koo, T. K. & Li, M. Y. A guideline of selecting and reporting intraclass correlation coefficients for reliability research. J. Chiropr. Med. 15, 155-163 (2016).
20. Benjamini, Y. & Hochberg, Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J. R. Stat. Soc. B 57, 289-300 (1995).
21. Cleveland, W. S. Robust locally weighted regression and smoothing scatterplots. J. Am. Stat. Assoc. 74, 829-836 (1979).
22. Chen, T. & Guestrin, C. XGBoost: a scalable tree boosting system. Proc. 22nd ACM SIGKDD Int. Conf. Knowl. Discov. Data Min. 785-794 (2016).
23. Ke, G. et al. LightGBM: a highly efficient gradient boosting decision tree. Adv. Neural Inf. Process. Syst. 30, 3149-3157 (2017).
24. Prokhorenkova, L. et al. CatBoost: unbiased boosting with categorical features. Adv. Neural Inf. Process. Syst. 31, 6638-6648 (2018).
25. Palinkas, L. A. et al. Purposeful sampling for qualitative data collection and analysis in mixed method implementation research. Adm. Policy Ment. Health 42, 533-544 (2015).
'''


DEEP_APPENDIX = r'''## S17. 审稿人追问级方法细节补充

### S17.1 120例病例构建的可复现算法

为避免“120例病例是否只是方便样本”的质疑，病例选择应在补充材料中写成可复现算法。具体流程为：

1. 建立候选池。以 OAI 中具备基线影像、症状/功能、BMI/体重、用药或合并症信息、至少一个治疗决策相关变量的病例为候选池。
2. 生成病例特征标签。每个候选病例生成结构严重度、症状负担、功能受限、代谢负荷、安全风险、治疗边界和管理需求标签。
3. 去除低信息病例。若病例缺少影像结构、症状/功能和核心安全字段中的任意两类以上，不进入标准化任务。
4. 去除重复场景。若多个病例在关键标签上高度相似，只保留信息最完整且源变量最清晰者。
5. 最大变异抽样。优先纳入能增加覆盖范围的病例，例如轻中重结构负担、肥胖/非肥胖、药物禁忌/无禁忌、高运动需求/低运动需求、手术边界明确/不明确、依从性风险高/低。
6. 锁定任务文本。将病例转为统一医生任务模板，并记录源变量、派生规则、缺失字段和版本号。
7. 覆盖性审计。导出覆盖矩阵，确认 120 例覆盖主要 KOA 管理场景，而不是某一类患者的重复。

该过程强调“任务代表性”而非“流行病学代表性”。论文中应避免写“随机代表真实世界患者分布”，而应写“覆盖医生端决策支持系统需要面对的主要病例形态”。

### S17.2 标准化病例模板

每个病例建议保存为如下结构。此结构既用于系统输入，也用于医生实验输入，确保人机比较的病例资料一致。

| 字段域 | 必填/可选 | 示例内容 | 用途 |
|---|---|---|---|
| 去标识化病例ID | 必填 | CASE-xxx | 任务追踪、去重和评分关联 |
| 患者基本信息 | 必填 | 年龄、性别、身高、体重、BMI | Profile 和风险模型输入 |
| 膝侧和病程 | 必填 | 左/右/双膝、症状持续时间 | 影像和症状匹配 |
| 症状负担 | 必填 | 疼痛、晨僵、活动受限、夜间或静息痛 | 治疗强度和安全门控 |
| 功能状态 | 必填 | KOOS/WOMAC、步行、上下楼、工作/运动需求 | 个体化处方 |
| 影像结构 | 必填 | KL、JSN、骨赘、硬化或影像模型输出 | KOM-Rad 和风险解释 |
| 合并症/用药风险 | 必填 | GI/renal/CV/抗凝/跌倒风险 | KOM-Safe |
| 治疗目标 | 必填 | 减痛、恢复步行、运动回归、延缓手术 | MDT 目标设定 |
| 既往治疗 | 可选 | 物理治疗、NSAIDs、注射、支具 | 方案升级和失败模式 |
| 随访边界 | 可选 | 复诊时间、升级条件、转诊边界 | KOM-Rx 输出 |

### S17.3 KOM-Profile 字段字典和评价分母

KOM-Profile 的 124 个字段不应全部进入同一个准确率分母。建议在补充表中分为三组：

| 字段组 | 数量 | 评价方式 | 是否进入主结果 |
|---|---:|---|---|
| Source-truth fields | 56 | exact/partial/wrong/missing | 是 |
| Rule-derived fields | 待导出 | 规则审计和人工 spot-check | 补充 |
| Management/free-text fields | 待导出 | 失败模式和可读性审计 | 补充 |

主文建议报告：“在 56 个可源真值核对字段中，字段级准确率/F1 为 0.846”。若需要报告 124 字段全框架，应写成“124 个字段框架，其中 56 个可核对字段进入准确率分母，68 个为规则派生或管理性字段”。这样能避免审稿人质疑 F1 分母混乱。

### S17.4 KOM-Rad/OAK-Net 训练细节模板

影像模块应在补充方法中记录如下训练配置。若某些配置未恢复，不应写成已完成，而应列为需要从训练日志恢复：

| 项目 | 建议记录 | 当前写作状态 |
|---|---|---|
| 数据来源 | OAI 膝关节 X 线影像和中心化读片标签 | 可写 |
| 标签 | KL 分级、JSN、骨赘、硬化等 | 可写 |
| 主模型 | ConvNeXt-B + evidential uncertainty head | 可写 |
| 对照模型 | DenseNet-121 | 可写 |
| Split | 患者级 split，同一患者影像不跨集合 | 可写原则 |
| 预处理 | 方向统一、ROI裁剪、灰度归一化、分辨率标准化 | 可写原则 |
| 增强 | 随机裁剪、旋转、亮度/对比度扰动 | 可写原则 |
| 优化器/学习率/epoch | 从训练日志恢复 | 若未找到则不写具体数 |
| 输出 | KL概率、预测等级、不确定性、医生复核提示 | 可写 |
| 指标 | QWK、BACC、macro-F1、MAE、ECE、sel_acc@80 | 可写 |

### S17.5 QM-Risk/KOM-Risk 输入指标清单

KOM-Risk 的“输入了哪些指标”建议以变量域而非逐个变量名写入主文，并把完整变量表放入补充表。当前可写的变量域包括：

| 变量域 | 代表变量 | 使用目的 |
|---|---|---|
| 人口学 | 年龄、性别、种族/中心如可用 | 基线风险调整 |
| 体格/代谢 | BMI、体重、身高 | 机械负荷和代谢风险 |
| 症状 | 疼痛、僵硬、KOOS/WOMAC pain | 症状恶化 endpoint 和风险预测 |
| 功能 | ADL、步行、上下楼、功能评分 | 功能恶化和治疗需求 |
| 影像结构 | KL、JSN、骨赘、硬化 | 结构进展风险 |
| 合并症/安全 | GI、renal、CV、跌倒、抗凝或相关用药风险 | 治疗安全门控 |
| 治疗史 | 既往 NSAIDs、注射、物理治疗、手术史 | 风险解释和临床分层 |
| 活动/行为 | 运动负荷、体力活动、依从性相关字段 | 管理需求和风险解释 |
| 缺失指示 | 关键字段缺失 indicator | 减少缺失机制偏倚 |

模型输出包括三个 endpoint 的连续概率、预设阈值下风险分层、校准状态、DCA 结果和解释输出。若阈值文件未恢复，论文只写“连续风险评分”和“预定义阈值待 locked threshold table 支持”，不要写具体低/中/高阈值。

### S17.6 KOM-Risk 训练和模型接受标准

每个 endpoint 的训练流程为：训练集拟合预处理器和候选模型；验证集选择模型和阈值；测试集进行一次性最终评价。候选模型包括 elastic-net logistic regression、random forest、XGBoost、LightGBM 和 CatBoost。模型接受标准如下：

| 标准 | 接受要求 |
|---|---|
| Endpoint 可复现 | 标签定义、随访窗口、事件率和纳排规则可复核 |
| 无泄漏 | 不使用未来影像、未来治疗、未来结局或 endpoint proxy |
| Person-level split | 同一受试者双膝不跨训练/验证/测试 |
| 判别能力 | AUROC/AUPRC 达到预设或相对基线有增益 |
| 校准 | Brier、校准曲线和 ECE 不显示严重失准 |
| 临床效用 | DCA 在合理阈值区间有净获益 |
| 解释合理性 | top features 符合 KOA 临床逻辑 |
| 可归档 | best_model、prediction、metrics、SHAP、QC 均保存 |

### S17.7 Evidence unit 数量一致性审计

当前上传草案明确写入 3,266 个 evidence units 和 2,174 个唯一来源。用户口头提到 3,666 个 evidence units，但在当前附件文本中未找到 3,666。论文写作应以可追溯 locked export 为准：

| 数字 | 当前证据 | 写作建议 |
|---|---|---|
| 3,266 evidence units | 上传草案 S6.2 明确写入 | 当前主文/补充方法采用此数 |
| 3,666 evidence units | 本轮附件未检出 | 仅在重新导出 evidence manifest 后替换 |
| 2,174 unique sources | 上传草案 S6.2 明确写入 | 可写 |
| 7,282 PubMed candidates | 2026-06-15 E-utilities 复核 | 可写为候选检索集 |

这四个数字代表不同层级：PubMed candidate records、unique sources、evidence units、evidence-level rows。写作中必须避免把候选文献数、唯一来源数和 evidence unit 数混用。

### S17.8 RAG gold standard 详细评价公式

对第 i 个 query，设返回列表为 R_i，金标准相关证据集合为 G_i，rel(r) 表示返回证据 r 的相关性等级。

Precision@10 = 前10条返回中 relevant evidence 数 / 10。  
Recall@K = 前K条返回命中的 gold evidence 数 / gold evidence 总数。  
Hit@10 = 若前10条至少包含1条 relevant evidence，则为1，否则为0。  
MRR = 1 / 第一个 relevant evidence 的排序位次；若无命中则为0。  
nDCG@10 = DCG@10 / IDCG@10，其中 DCG 使用等级相关性折损排序位置。

若 gold label 含 partially relevant，可在主分析中只把 relevant 视为阳性，在敏感性分析中把 partial relevant 赋予 0.5 权重。该权重规则必须预先定义。

### S17.9 MDT prompt 模板

每个 MDT 智能体建议使用固定 prompt 模板：

```text
Role: [specialty agent role]
Inputs: standardized patient profile, available imaging/risk outputs, retrieved evidence units, safety flags
Task: produce specialty-specific recommendations only within role boundary
Evidence rule: cite evidence unit IDs; do not invent citations; downgrade unsupported suggestions
Safety rule: list contraindications, stop rules, referral boundaries and missing information
Output format:
1. Key interpretation
2. Recommended action
3. Patient-specific adjustment
4. Evidence IDs
5. Safety constraints
6. Uncertainty / doctor review required
```

R5 跨专科质疑 prompt 应强制寻找冲突、遗漏和禁忌。R6 证据仲裁 prompt 应强制按 L1-L7 和 direct/context/safety evidence 进行保留、降级或删除。R8 最终处方 prompt 不允许新增未被 R1-R7 支持的治疗建议。

### S17.10 KOM-Sim 统计模型细化

KOM-Sim 的主要分析单元为医生-任务级处方记录。推荐主模型为混合效应模型：

```text
Score_ijk = beta0 + beta1*Condition_j + beta2*CaseCovariates_k
            + u_physician_i + u_case_k + error_ijk
```

其中 u_physician 表示医生随机效应，u_case 表示病例随机效应。若评分为连续近似正态，可用线性混合模型；若为有序等级评分，可用 cumulative-link mixed model；若为安全关键错误，可用 mixed-effects logistic 或 Poisson/negative-binomial model。主比较为 Clinician + KOM vs Clinician alone；次比较为 Clinician + KOM-R vs Clinician + KOM。所有比较需报告效应量、95% CI 和 FDR 校正后的 p 值。

### S17.11 消融实验评价矩阵

KOM 消融实验应统一使用同一病例输入、同一处方格式、同一评分者和同一 KOM-Score。评价指标包括：

| 维度 | 指标 | 解释 |
|---|---|---|
| 质量 | overall expert score | 处方整体临床质量 |
| 安全 | safety expert score, safety-critical errors | 是否遗漏关键安全门控 |
| 指南一致性 | guideline alignment | 是否符合 L1/L2 证据和指南锚点 |
| 个体化 | personalization | 是否根据 BMI、症状、功能、禁忌和目标调整 |
| 可执行性 | actionability | 是否给出可实施的运动、药物、随访参数 |
| 证据可追溯性 | evidence traceability | 是否关联 evidence unit ID |
| 完整性 | rule score | 是否覆盖核心治疗组件 |

Full KOM 相对 w/o RAG 的差异解释证据检索贡献；Full KOM 相对 w/o MDT 的差异解释多专科协商贡献；Full KOM 相对 Direct LLM 的差异解释结构化专科系统整体贡献。

'''


CHECKLIST = r'''# KOM 补充方法学覆盖性自查表（2026-06-15）

| 用户要求模块 | 已写入位置 | 状态 | 说明 |
|---|---|---|---|
| 120例标准化病例如何纳入 | S2.1-S2.5 | 已补充 | 改为多维覆盖矩阵，不再使用四象限作为论文框架 |
| 如何证明病例选择符合原设计 | S2.5 | 已补充 | 加入 manifest、去重、完整性、多维覆盖、版本冻结审计 |
| 120例不是同一个患者、代表多样患者 | S2.1-S2.4 | 已补充 | 强调去标识化患者单位和最大变异抽样 |
| KOM-Profile 自动抽取患者画像 | S3 | 已补充 | 输入、字段分层、规则化构建、评价指标和锁定结果 |
| 影像分析输入、输出、架构、实验设计 | S4 | 已补充 | KOM-Rad/OAK-Net、ConvNeXt-B、DenseNet、指标和结果 |
| QM-Risk/KOM-Risk 深度更新 | S5 | 已补充 | endpoints、特征数、去重、split、候选算法、CatBoost locked result |
| SHAP 分析 | S6 | 已补充 | 全局/个体解释及非因果边界 |
| 证据库如何建立 | S7 | 已补充 | PubMed 检索式、候选 count、evidence unit 字段和层级 |
| PubMed 检索词与候选量验证 | S7.1 | 已补充 | 2026-06-15 E-utilities count = 7,282 |
| L1-L7 证据层级 | S7.3 | 已补充/带边界 | 定义 L1-L7；现有 locked count 为 L1-L5 压缩层级，未硬拆 |
| RAG 返回流程 | S8.1-S8.2 | 已补充 | query、路由、direct/context/safety evidence 和 baseline |
| RAG gold standard 与指标计算 | S8.3-S8.4 | 已补充 | 写入 gold label 人工审阅逻辑和 Precision/Recall/Hit/MRR/nDCG |
| RAG 测试结果 | S8.5 | 已补充 | 写入 160 query 与锁定指标 |
| MDT 设计和每个智能体构建 | S9 | 已补充 | R0-R8 表格、prompt 构成、知识源和交互流程 |
| KVMSafe/KOM-Safe 审计标准 | S10 | 已补充 | 安全域、pass/warning/critical、与智能体交互 |
| KM-score/KOM-Score 是什么 | S11 | 已补充 | Expert、Rule、Error 三部分及公式 |
| KVMSIM/KOM-Sim 医生交互实验 | S12 | 已补充 | 26名医生、30任务、780记录、条件和记录变量 |
| 消融实验 | S13 | 已补充 | Full KOM、w/o RAG、w/o MDT、Direct LLM |
| 图件位置预留 | S1-S16 | 已补充 | 图S1-S9 placeholders |
| 不可写成已完成的缺口 | S15 | 已补充 | RAG node/edge、generation grounding、OAK-Net image-level prediction 等 |
'''


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)


def add_table_from_rows(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    total = 9360
    col_width = total // len(rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, col_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = ""
            cell.text = value.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)
            if r_idx == 0:
                set_cell_shading(cell, "F4F6F9")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
    doc.add_paragraph()


def markdown_to_docx(md, path):
    doc = Document()
    style_doc(doc)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                if set(tl.replace("|", "").strip()) <= {"-", ":"}:
                    continue
                rows.append([c.strip() for c in tl.strip("|").split("|")])
            add_table_from_rows(doc, rows)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.1)
            p.add_run(stripped[2:])
        elif stripped.startswith("[图 "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped)
            run.italic = True
            run.font.color.rgb = RGBColor(31, 77, 120)
        elif stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(stripped)
        i += 1

    doc.save(path)


def main():
    full_content = CONTENT.replace("## S17. 参考文献和方法学依据", DEEP_APPENDIX + "\n## S22. 参考文献和方法学依据")
    MD_PATH.write_text(full_content, encoding="utf-8")
    CHECK_PATH.write_text(CHECKLIST, encoding="utf-8")
    markdown_to_docx(full_content, DOCX_PATH)
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {CHECK_PATH}")


if __name__ == "__main__":
    main()
